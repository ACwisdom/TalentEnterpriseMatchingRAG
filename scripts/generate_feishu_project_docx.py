"""
生成根目录「飞书用」项目说明 Word（.docx），便于上传飞书知识库或发附件。

用法（仓库根目录）:
  python scripts/generate_feishu_project_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    # 根目录文件名使用 ASCII，避免部分 Windows 终端/工具链编码导致乱码；正文仍为中文。
    out = root / "TalentEnterpriseMatchingRAG_project_brief_feishu.docx"

    doc = Document()
    t = doc.add_heading("TalentEnterpriseMatchingRAG 项目说明", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    r = sub.add_run("（适用于飞书文档 / 知识库：可直接上传本 .docx 或复制正文）")
    r.italic = True
    r.font.size = Pt(10)

    _h(doc, "一、项目定位", 1)
    _p(
        doc,
        "人才—企业智能匹配与报告生成系统（RAG）。以企业 Excel 库为数据源，构建向量索引，"
        "融合关键词、向量与混合检索（如 RRF），从人才简历中召回匹配企业；可接入小米 MiMo 生成匹配报告与人才画像，"
        "并导出 Word。提供 Web 上传、流式展示、Demo 目录与部署方案，适用于引才咨询、政企对接中的初筛与材料起草辅助。",
    )

    _h(doc, "二、核心功能", 1)
    rows = [
        ("企业检索", "关键词 / 向量 / 混合，配置见 src/config.py 中 RETRIEVAL_MODE"),
        ("向量索引", "Chroma + sentence-transformers，持久化目录 data/enterprise_vectors"),
        ("简历解析", "PDF、DOCX、TXT（src/resume_processor.py）"),
        ("大模型", "MiMo 生成匹配报告、人才画像；可选高级画像（见 docs/高级人才画像功能说明.md）"),
        ("匹配与重排", "宽召回、CrossEncoder 与规则重排、可选 LTR；详见 docs/match-rerank.md"),
        ("Web", "FastAPI：流式 POST /api/match/stream、JSON POST /api/match、/demo、/health（web/main.py）"),
        ("Java 扩展 API", "java/recruitment-api：企业/职位/人才/推荐/沟通/消息/提醒 REST v1，OpenAPI / Swagger"),
        ("LangGraph 工具层", "langgraph_tools：httpx 调用 Java，示例图 example_graph.py；本阶段不做 MCP"),
        ("部署", "Docker、fly.toml；说明见 docs/deploy-public.md"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b

    _h(doc, "三、技术栈（摘要）", 1)
    _p(
        doc,
        "Python 3.11+；FastAPI、uvicorn；langchain；Chroma；scikit-learn / sentence-transformers（视检索模式）；"
        "可选 Java 21 + Spring Boot 3（独立子目录，可迁出单独仓库）；LangGraph + httpx 用于编排侧调用 Java。",
    )

    _h(doc, "四、目录结构（节选）", 1)
    _p(
        doc,
        "src/ — 检索、流水线、LLM、Word 导出等核心逻辑\n"
        "web/ — FastAPI 与静态页\n"
        "java/recruitment-api/ — 招聘域 REST v1\n"
        "langgraph_tools/ — Java HTTP 客户端与 LangGraph 示例\n"
        "contracts/ — OpenAPI 契约快照\n"
        "data/ — 企业表、向量库、输入输出（业务数据勿随意提交）\n"
        "docs/ — 部署、画像、项目介绍、Java 联调等\n"
        "scripts/ — 批处理、训练与辅助脚本\n"
        "prompts/ — 系统提示词",
    )

    _h(doc, "五、环境与关键变量", 1)
    _p(doc, "Python 侧（.env 或环境变量，按需配置）：", bold=True)
    vars_py = [
        "MIMO_API_KEY — 调用 MiMo（可选）",
        "APP_API_KEY — Web 接口 X-API-Key（公网访问 /api/match* 建议配置）",
        "MIMO_USE_STREAM — 是否使用流式（可关闭回退同步）",
        "RERANK_ENABLED 等 — 见 src/config.py 与 docs/match-rerank.md",
        "JAVA_API_BASE_URL / JAVA_SERVICE_API_KEY — 调用 Java 时配置（见 docs/java-langgraph-integration.md）",
    ]
    for line in vars_py:
        doc.add_paragraph(line, style="List Bullet")

    _p(doc, "Windows 建议：PYTHONUTF8=1 或 PYTHONIOENCODING=utf-8，避免控制台编码导致 Web 异常。", bold=True)

    _h(doc, "六、快速开始", 1)
    _p(doc, "1. 创建虚拟环境并 pip install -r requirements.txt", bold=True)
    _p(doc, "2. 准备 data/enterprise_db.xlsx，执行：python -m src.build_enterprise_index")
    _p(doc, "3. 本地 Web：uvicorn web.main:app --host 127.0.0.1 --port 8000，浏览器访问 http://127.0.0.1:8000/")
    _p(doc, "4. 测试：pytest tests/")
    _p(doc, "5. Java 子项目：cd java/recruitment-api && mvn spring-boot:run（需 JDK 21 + Maven）")

    _h(doc, "七、流式与反向代理注意", 1)
    _p(
        doc,
        "首页使用 SSE（text/event-stream）。若经 Nginx/Caddy 等反向代理，请关闭响应缓冲并酌情调大读超时，"
        "避免流式被攒包或中途断开。细节见 docs/deploy-public.md。",
    )

    _h(doc, "八、文档索引（仓库内 Markdown）", 1)
    links = [
        "README.md — 总览与命令",
        "docs/项目介绍.md — 约 300 字简介",
        "docs/deploy-public.md — 公网部署",
        "docs/高级人才画像功能说明.md",
        "docs/match-rerank.md — 匹配与重排",
        "docs/java-langgraph-integration.md — Java 与 LangGraph 联调",
    ]
    for x in links:
        doc.add_paragraph(x, style="List Bullet")

    _h(doc, "九、合规与数据", 1)
    _p(
        doc,
        "业务数据与向量目录通常在 data/ 下；协作时请勿将含隐私的生产数据提交到公开仓库。"
        "许可证：仓库未统一声明 LICENSE 前，默认保留所有权利。",
    )

    _h(doc, "十、文档维护", 1)
    _p(
        doc,
        "本 Word 由 scripts/generate_feishu_project_docx.py 生成。"
        "更新仓库说明后，可重新运行脚本覆盖根目录 TalentEnterpriseMatchingRAG_project_brief_feishu.docx。",
    )

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
