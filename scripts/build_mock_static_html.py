"""
从固定路径的 docx 生成 data/default/mock_full_page.html（纯静态，无需后端）。
路径写死在下方，可改后重新运行：python scripts/build_mock_static_html.py
"""

from __future__ import annotations

import base64
import html
from pathlib import Path

from docx import Document

PROJECT = Path(__file__).resolve().parent.parent
OUT_HTML = PROJECT / "data" / "default" / "mock_full_page.html"

RESUME_DOCX = PROJECT / "data" / "default" / "input" / "WZ-00003105王明华博士简历.docx"
PREVIEW_DOCX = Path(r"C:\Users\Lenovo\Desktop\报告预览.docx")
DOWNLOAD_DOCX = Path(
    r"C:\Users\Lenovo\Desktop\WZ-00003105王明华博士简历_匹配报告_20260510_213646_409787.docx"
)
DOWNLOAD_NAME = "WZ-00003105王明华博士简历_匹配报告_20260510_213646_409787.docx"


def docx_to_plain(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def main() -> None:
    if not RESUME_DOCX.is_file():
        raise SystemExit(f"缺少简历: {RESUME_DOCX}")
    if not PREVIEW_DOCX.is_file():
        raise SystemExit(f"缺少报告预览: {PREVIEW_DOCX}")
    if not DOWNLOAD_DOCX.is_file():
        raise SystemExit(f"缺少下载用 Word: {DOWNLOAD_DOCX}")

    resume_text = docx_to_plain(RESUME_DOCX)
    preview_text = docx_to_plain(PREVIEW_DOCX)
    docx_b64 = base64.b64encode(DOWNLOAD_DOCX.read_bytes()).decode("ascii")
    data_href = (
        "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,"
        + docx_b64
    )

    resume_esc = html.escape(resume_text, quote=True)
    preview_esc = html.escape(preview_text, quote=True)
    fname_esc = html.escape(DOWNLOAD_NAME, quote=True)

    body = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>人才-企业匹配（静态演示）</title>
  <style>
    :root {{ font-family: system-ui, sans-serif; color: #1a1a1a; background: #f4f5f7; }}
    body {{ max-width: 52rem; margin: 0 auto; padding: 1.5rem; }}
    h1 {{ font-size: 1.35rem; font-weight: 600; }}
    .banner {{ background: #fef3c7; border: 1px solid #f59e0b; color: #92400e; padding: 0.6rem 0.75rem; border-radius: 6px; font-size: 0.85rem; margin-bottom: 1rem; }}
    fieldset {{ border: 1px solid #cfd3dc; border-radius: 8px; padding: 1rem 1.1rem; margin: 0 0 1rem; background: #fff; }}
    legend {{ padding: 0 0.35rem; font-weight: 600; }}
    label {{ display: block; margin: 0.5rem 0 0.2rem; font-size: 0.9rem; }}
    input[type="text"], input[type="number"], input[type="password"] {{
      width: 100%; max-width: 22rem; padding: 0.45rem 0.5rem; border: 1px solid #bbb; border-radius: 4px; box-sizing: border-box;
    }}
    textarea.fake-file {{ width: 100%; max-width: 100%; min-height: 10rem; font-size: 0.82rem; padding: 0.5rem; border: 1px solid #bbb; border-radius: 4px; box-sizing: border-box; }}
    .row {{ display: flex; flex-wrap: wrap; gap: 1rem; align-items: center; margin-top: 0.5rem; }}
    .row label {{ display: flex; align-items: center; gap: 0.35rem; margin: 0; font-weight: normal; }}
    button {{ margin-top: 0.75rem; padding: 0.55rem 1.1rem; border-radius: 6px; border: none; background: #94a3b8; color: #fff; font-weight: 600; cursor: not-allowed; }}
    #status {{ margin: 0.75rem 0; font-size: 0.9rem; color: #444; min-height: 1.2em; }}
    #report {{ white-space: pre-wrap; background: #fff; border: 1px solid #cfd3dc; border-radius: 8px; padding: 1rem; max-height: 28rem; overflow: auto; font-size: 0.85rem; line-height: 1.45; }}
    #dl a {{ color: #2563eb; font-weight: 600; }}
    small {{ color: #666; }}
    .fname {{ font-size: 0.85rem; color: #334155; margin: 0.25rem 0 0.35rem; }}
  </style>
</head>
<body>
  <h1>人才-企业匹配</h1>
  <div class="banner">本页为<strong>静态演示</strong>：不连接服务器、不调用接口。可直接用浏览器打开本文件（file:// 或双击）。</div>

  <form id="f" onsubmit="return false;">
    <fieldset>
      <legend>请求（演示）</legend>
      <div id="apiKeyRow" style="display:none">
        <label for="apiKey">X-API-Key</label>
        <input id="apiKey" type="password" autocomplete="off" value="" />
      </div>

      <label>简历文件（已固定为项目内 docx 的正文提取）</label>
      <p class="fname">WZ-00003105王明华博士简历.docx</p>
      <textarea class="fake-file" readonly id="resumePreview">{resume_esc}</textarea>

      <label for="topK">top_k</label>
      <input id="topK" type="number" min="1" max="200" value="15" readonly />

      <div class="row">
        <label><input type="checkbox" id="useLlm" disabled /> use_llm（演示已关闭）</label>
        <label><input type="checkbox" id="useLlmProfile" disabled /> use_llm_for_profile（演示已关闭）</label>
      </div>

      <button type="button" id="go" disabled>提交匹配（演示无效）</button>
    </fieldset>
  </form>

  <div id="status">完成。（静态演示）</div>
  <section id="out">
    <h2 style="font-size:1.05rem">报告预览</h2>
    <pre id="report">{preview_esc}</pre>
    <div id="dl">
      <p><a id="wordDl" href="{data_href}" download="{fname_esc}">下载 Word：{fname_esc}</a></p>
      <p><small>内嵌为 base64，体积约 {len(docx_b64)} 字符；若浏览器拦截下载，请右键链接另存为。</small></p>
    </div>
  </section>
</body>
</html>
"""
    OUT_HTML.parent.mkdir(parents=True, exist_ok=True)
    OUT_HTML.write_text(body, encoding="utf-8")
    print(f"Wrote {OUT_HTML}")


if __name__ == "__main__":
    main()
