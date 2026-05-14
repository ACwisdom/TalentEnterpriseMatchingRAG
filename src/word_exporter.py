"""
Word文档导出模块
功能：将匹配报告导出为格式化的Word文档
"""
import os
import sys
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx未安装，Word导出功能不可用")
    print("   请运行: pip install python-docx")
    DOCX_AVAILABLE = False

# 添加项目根目录到路径
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


class WordExporter:
    """Word文档导出器"""
    
    def __init__(self):
        """初始化Word导出器"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx未安装，无法使用Word导出功能")
        
        print("✅ Word导出器初始化成功")
    
    def create_matching_report(self,
                               resume_text: str,
                               enterprises: List[Dict],
                               output_path: str,
                               report_title: str = "人才-企业匹配报告",
                               talent_profile: str = None,
                               is_advanced_profile: bool = False) -> str:
        """
        创建匹配报告Word文档
        
        Args:
            resume_text: 简历文本
            enterprises: 企业列表
            output_path: 输出文件路径
            report_title: 报告标题
            talent_profile: 人才画像摘要（由LLM生成），如果为None则使用截取文本
            is_advanced_profile: 是否为高级人才画像（结合简历+微信截图）
            
        Returns:
            str: 保存的文件路径
        """
        print(f"\n📝 正在创建Word文档: {output_path}")
        
        # 创建文档
        doc = Document()
        
        # 设置文档默认字体（中文支持）
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10.5)
        
        # ============ 标题 ============
        title = doc.add_heading(report_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        time_run = time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_run.font.size = Pt(9)
        time_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # 空行
        
        # ============ 人才画像摘要 ============
        if is_advanced_profile:
            doc.add_heading('📋 高级人才画像（结合简历+微信分析）', level=1)
        else:
            doc.add_heading('📋 人才画像摘要', level=1)
        
        # 如果有LLM生成的人才画像，使用它；否则使用截取文本
        if talent_profile:
            # 根据是否为高级人才画像，使用不同的格式化方法
            if is_advanced_profile:
                self._add_advanced_profile(doc, talent_profile)
            else:
                p = doc.add_paragraph()
                run = p.add_run(talent_profile)
                run.font.size = Pt(10)
        else:
            # 简历文本（截取前1000字符）
            resume_summary = resume_text[:1000] + "..." if len(resume_text) > 1000 else resume_text
            p = doc.add_paragraph()
            run = p.add_run(resume_summary)
            run.font.size = Pt(10)
        
        doc.add_paragraph()  # 空行
        
        # ============ 企业匹配结果 ============
        doc.add_heading('🏢 企业匹配结果', level=1)
        
        # 统计信息
        stats_para = doc.add_paragraph()
        stats_run = stats_para.add_run(f"检索概况：共匹配 {len(enterprises)} 家企业\n")
        stats_run.font.bold = True
        
        stats_para.add_run("数据来源：附件《沃咨企业库.xlsx》")
        
        doc.add_paragraph()  # 空行
        
        # 按匹配度分类
        high_match = [e for e in enterprises if e.get('score', 0) >= 70]
        potential_match = [e for e in enterprises if 40 <= e.get('score', 0) < 70]
        related = [e for e in enterprises if e.get('score', 0) < 40]
        
        # 高度匹配企业
        if high_match:
            doc.add_heading('✅ 高度匹配企业（优先推荐）', level=2)
            self._add_enterprise_table(doc, high_match, "高度匹配")
        
        # 潜在匹配企业
        if potential_match:
            doc.add_heading('⚠️ 潜在匹配企业（可考虑）', level=2)
            self._add_enterprise_table(doc, potential_match, "潜在匹配")
        
        # 行业相关企业
        if related:
            doc.add_heading('🔶 行业相关企业（拓展参考）', level=2)
            self._add_enterprise_table(doc, related, "行业相关")
        
        # ============ 匹配总结 ============
        doc.add_page_break()
        doc.add_heading('📊 匹配总结', level=1)
        
        # 创建统计表格
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        table.rows[0].cells[0].text = "指标"
        table.rows[0].cells[1].text = "数值"
        
        # 数据行
        table.rows[1].cells[0].text = "高度匹配"
        table.rows[1].cells[1].text = f"{len(high_match)} 家"
        
        table.rows[2].cells[0].text = "潜在匹配"
        table.rows[2].cells[1].text = f"{len(potential_match)} 家"
        
        table.rows[3].cells[0].text = "行业相关"
        table.rows[3].cells[1].text = f"{len(related)} 家"
        
        table.rows[4].cells[0].text = "合计"
        table.rows[4].cells[1].text = f"{len(enterprises)} 家"
        
        # 添加备注
        doc.add_paragraph()
        note_para = doc.add_paragraph()
        note_para.add_run("⚠️ 以上企业信息均来源于附件《沃咨企业库.xlsx》，建议在正式接触前进一步核实最新情况。")
        note_para.add_run("\n注：附件中未提供的字段已标注'附件中未提供'，未编造任何信息。")
        
        # ============ 保存文档 ============
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        print(f"✅ Word文档已保存: {output_path}")
        print(f"   文件大小: {os.path.getsize(output_path)} 字节")
        
        return output_path
    
    def _add_enterprise_table(self, doc: Document, enterprises: List[Dict], category: str):
        """
        添加企业信息表格
        
        Args:
            doc: Word文档对象
            enterprises: 企业列表
            category: 分类名称
        """
        for i, e in enumerate(enterprises, 1):
            metadata = e.get('metadata', {})
            
            # 企业名称（作为标题）
            p = doc.add_paragraph()
            run = p.add_run(f"{i}、{metadata.get('企业名称', '（未知）')}")
            run.font.size = Pt(12)
            run.font.bold = True
            
            # 创建信息表格
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Light List Accent 1'
            
            # 填充表格数据
            table.rows[0].cells[0].text = "匹配度评分"
            table.rows[0].cells[1].text = f"{e.get('score', 0):.2f}/100"
            
            table.rows[1].cells[0].text = "所在地区"
            region = f"{metadata.get('省', '')}{metadata.get('市', '')}{metadata.get('区/县', '')}"
            table.rows[1].cells[1].text = region if region else "附件中未提供"
            
            table.rows[2].cells[0].text = "地区门槛"
            table.rows[2].cells[1].text = metadata.get('地区门槛', '附件中未提供')
            
            table.rows[3].cells[0].text = "一级领域"
            table.rows[3].cells[1].text = metadata.get('一级领域', '附件中未提供')
            
            table.rows[4].cells[0].text = "二级领域"
            table.rows[4].cells[1].text = metadata.get('二级领域', '附件中未提供')
            
            table.rows[5].cells[0].text = "企业主要产品"
            table.rows[5].cells[1].text = metadata.get('企业主要产品', '附件中未提供')
            
            # 匹配关键词
            if e.get('matched_keywords'):
                doc.add_paragraph(f"匹配关键词: {', '.join(e['matched_keywords'][:8])}")
            
            doc.add_paragraph()  # 空行
        
        doc.add_paragraph()  # 分类之间加空行
    
    def _add_advanced_profile(self, doc: Document, profile_text: str):
        """
        格式化高级人才画像（结合简历+微信分析）
        
        Args:
            doc: Word文档对象
            profile_text: 高级人才画像文本（Markdown格式）
        """
        # 按行分割
        lines = profile_text.split('\n')
        
        current_heading_level = None
        current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # 检测Markdown标题
            if line_stripped.startswith('## '):
                # 保存上一段内容
                if current_heading_level is not None:
                    self._add_formatted_section(doc, current_heading_level, current_content)
                
                # 新标题
                current_heading_level = 2  # Word标题级别
                current_content = [line_stripped[3:]]  # 去掉"## "
                
            elif line_stripped.startswith('### '):
                # 保存上一段内容
                if current_heading_level is not None:
                    self._add_formatted_section(doc, current_heading_level, current_content)
                
                # 新标题
                current_heading_level = 3  # Word标题级别
                current_content = [line_stripped[4:]]  # 去掉"### "
                
            elif line_stripped.startswith('#### '):
                # 保存上一段内容
                if current_heading_level is not None:
                    self._add_formatted_section(doc, current_heading_level, current_content)
                
                # 新标题
                current_heading_level = 4  # Word标题级别
                current_content = [line_stripped[5:]]  # 去掉"#### "
                
            else:
                # 普通内容行
                if line_stripped:  # 非空行
                    current_content.append(line_stripped)
                else:
                    # 空行，保存上一段内容
                    if current_heading_level is not None and current_content:
                        self._add_formatted_section(doc, current_heading_level, current_content)
                        current_content = []
        
        # 保存最后一段内容
        if current_heading_level is not None and current_content:
            self._add_formatted_section(doc, current_heading_level, current_content)
    
    def _add_formatted_section(self, doc: Document, heading_level: int, content_lines: List[str]):
        """
        将格式化的段落添加到Word文档
        
        Args:
            doc: Word文档对象
            heading_level: 标题级别（1-4）
            content_lines: 内容行列表（第一行是标题）
        """
        if not content_lines:
            return
        
        # 添加标题（第一行）
        title = content_lines[0]
        doc.add_heading(title, level=heading_level)
        
        # 添加内容（剩余行）
        if len(content_lines) > 1:
            content = '\n'.join(content_lines[1:])
            
            # 检测是否为列表格式
            if content.startswith('- ') or content.startswith('* ') or '• ' in content:
                # 列表格式，逐行添加
                for line in content.split('\n'):
                    line = line.strip()
                    if line.startswith('- ') or line.startswith('* '):
                        # 列表项
                        p = doc.add_paragraph(line[2:], style='List Bullet')
                        p.paragraph_format.space_after = Pt(2)
                    elif line.startswith('  - ') or line.startswith('  * '):
                        # 子列表项
                        p = doc.add_paragraph(line.strip()[2:], style='List Bullet 2')
                        p.paragraph_format.space_after = Pt(2)
                    elif line:
                        # 普通文本
                        p = doc.add_paragraph(line)
                        p.paragraph_format.space_after = Pt(2)
            else:
                # 普通段落
                p = doc.add_paragraph(content)
                p.paragraph_format.space_after = Pt(4)
        
        doc.add_paragraph()  # 空行分隔
    
    def export_pipeline_result(self,
                              resume_text: str,
                              enterprises: List[Dict],
                              output_dir: str = None,
                              talent_profile: str = None) -> str:
        """
        导出Pipeline结果为Word文档（便捷函数）
        
        Args:
            resume_text: 简历文本
            enterprises: 企业列表
            output_dir: 输出目录，默认使用config中的OUTPUT_DIR
            talent_profile: 人才画像摘要（由LLM生成）
            
        Returns:
            str: 保存的文件路径
        """
        if output_dir is None:
            from src.config import OUTPUT_DIR
            output_dir = OUTPUT_DIR
        
        # 生成文件名（带时间戳）
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"人才企业匹配报告_{timestamp}.docx"
        output_path = os.path.join(output_dir, filename)
        
        return self.create_matching_report(resume_text, enterprises, output_path, talent_profile=talent_profile)


def export_to_word(resume_text: str,
                   enterprises: List[Dict],
                   output_path: str) -> str:
    """
    快速导出为Word的便捷函数
    
    Args:
        resume_text: 简历文本
        enterprises: 企业列表
        output_path: 输出文件路径
        
    Returns:
        str: 保存的文件路径
    """
    exporter = WordExporter()
    return exporter.create_matching_report(resume_text, enterprises, output_path)


if __name__ == "__main__":
    # 测试代码
    print("=" * 60)
    print("测试Word导出模块")
    print("=" * 60)
    
    try:
        # 1. 初始化导出器
        exporter = WordExporter()
        
        # 2. 创建测试数据
        test_resume = """
        张三，男，1985年出生，博士学历，专业方向：生物医药、药物研发。
        教育背景：
        - 2008-2012：XX大学 药学 学士
        - 2012-2015：XX大学 药物化学 硕士
        - 2015-2019：XX大学 药学 博士
        
        工作经历：
        - 2019-至今：XX制药企业 高级研究员
          - 负责化学制药工艺开发
          - 参与多个药物研发项目
          - 发表SCI论文10篇
        
        专业技能：
        - 药物研发全流程经验
        - 化学制药工艺优化
        - 生物医药技术转化
        - 项目管理经验丰富
        """
        
        test_enterprises = [
            {
                "metadata": {
                    "企业名称": "测试医药企业A",
                    "省": "江苏省",
                    "市": "苏州市",
                    "区/县": "工业园区",
                    "地区门槛": "博士可享安家补贴30万",
                    "一级领域": "4. 生物医药技术",
                    "二级领域": "4.1 现代医药 - 化学",
                    "企业主要产品": "化学原料药、制剂、创新药研发"
                },
                "score": 85.5,
                "matched_keywords": ["药物研发", "化学制药", "生物医药", "博士"]
            },
            {
                "metadata": {
                    "企业名称": "测试生物科技企业B",
                    "省": "江苏省",
                    "市": "南京市",
                    "区/县": "栖霞区",
                    "地区门槛": "附件中未提供",
                    "一级领域": "4. 生物医药技术",
                    "二级领域": "4.2 生物制药",
                    "企业主要产品": "单克隆抗体、重组蛋白药物"
                },
                "score": 72.3,
                "matched_keywords": ["生物医药", "博士", "研发"]
            }
        ]
        
        # 3. 导出Word文档
        output_file = "data/output/test_report.docx"
        exporter.create_matching_report(
            resume_text=test_resume,
            enterprises=test_enterprises,
            output_path=output_file
        )
        
        print("\n" + "=" * 60)
        print("✅ Word导出测试成功！")
        print("=" * 60)
        print(f"输出文件: {output_file}")
        print(f"请打开文件检查格式是否正确")
        
    except Exception as e:
        print(f"\n❌ 测试失败: {str(e)}")
        import traceback
        traceback.print_exc()
