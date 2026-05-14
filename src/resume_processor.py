"""
简历处理模块
功能：读取多种格式的简历文件（PDF、DOCX、TXT）
"""
import os
import sys
from pathlib import Path
from typing import TYPE_CHECKING, List, Optional

# 添加项目根目录到路径
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from src.pipeline import run_pipeline_and_export_word
from src.config import OUTPUT_DIR

if TYPE_CHECKING:
    from src.match_constraints import MatchConstraints


def _collapse_merged_cell_row_texts(cell_texts: list[str]) -> list[str]:
    """
    Word 合并单元格在 python-docx 中常表现为同一文本在行内重复多格。
    折叠为「逻辑列」：若本行所有非空单元格全文相同则只保留一条；否则顺序去掉连续重复。
    """
    stripped = [(t or "").strip() for t in cell_texts]
    non_empty = [t for t in stripped if t]
    if not non_empty:
        return []
    if len(non_empty) > 1 and len(set(non_empty)) == 1:
        return [non_empty[0]]
    out: list[str] = []
    for t in stripped:
        if not t:
            continue
        if out and out[-1] == t:
            continue
        out.append(t)
    return out


def read_resume_file(file_path: str) -> str:
    """
    读取简历文件，支持PDF、DOCX、TXT格式
    
    Args:
        file_path: 简历文件路径
        
    Returns:
        str: 简历纯文本内容
        
    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 不支持的文件格式
        Exception: 读取失败
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ 简历文件不存在: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    print(f"📂 正在读取简历文件: {file_path}")
    print(f"   文件格式: {file_ext}")
    
    try:
        if file_ext == '.txt':
            # 读取纯文本文件
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"✅ TXT文件读取成功（{len(content)} 字符）")
            return content
        
        elif file_ext == '.docx':
            # 读取Word文档
            try:
                from docx import Document
            except ImportError:
                raise ImportError("python-docx未安装，无法读取DOCX文件。请运行: pip install python-docx")
            
            doc = Document(file_path)
            content_parts = []
            
            # 读取所有段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())
            
            # 读取所有表格内容（智能解析表格结构）
            for table in doc.tables:
                content_parts.append("")  # 与正文区分的空行，不插入 [表格 N] 占位符
                
                # 使用状态机解析表格
                state = "LOOKING_FOR_TITLE"
                headers = []
                
                for row_idx, row in enumerate(table.rows):
                    # 获取本行所有单元格（保留空单元格以维持索引对应）
                    all_cells = [cell.text.strip() for cell in row.cells]
                    non_empty = _collapse_merged_cell_row_texts(all_cells)
                    
                    if not non_empty:
                        continue  # 跳过空行
                    
                    # 去重后判断（处理合并单元格导致的重复）
                    unique_texts = list(dict.fromkeys(non_empty))
                    
                    # 判断是否是标题行（包含 、且较短，且去重后只有1个唯一值）
                    if len(unique_texts) == 1:
                        text = unique_texts[0]
                        if ('、' in text or '（' in text) and len(text) < 30:
                            # 标题行
                            content_parts.append(f'\n{text}')
                            if text.startswith(('一、', '二、', '三、', '四、')):
                                if text.startswith('一、'):
                                    state = "EXPECT_HEADER"
                                else:
                                    state = "EXPECT_CONTENT"
                            continue
                    
                    # 根据状态处理
                    if state == "EXPECT_HEADER":
                        # 表头行
                        headers = list(non_empty)
                        state = "EXPECT_DATA"
                    elif state == "EXPECT_DATA" and headers:
                        # 数据行：与表头配对（用已折叠的 non_empty 与表头按索引对齐）
                        for i, header in enumerate(headers):
                            if i < len(non_empty):
                                content_parts.append(f"{header}: {non_empty[i]}")
                        headers = []
                        state = "LOOKING_FOR_TITLE"
                    elif state == "EXPECT_CONTENT":
                        # 内容行（合并单元格已折叠，不再把重复格子拼成多行）
                        content = non_empty[0] if len(non_empty) == 1 else "\n".join(non_empty)
                        content_parts.append(content)
                        state = "LOOKING_FOR_TITLE"
                    else:
                        # 未预期的状态，按普通内容处理
                        content = non_empty[0] if len(non_empty) == 1 else ' | '.join(non_empty)
                        content_parts.append(content)
            
            # 合并内容
            content = '\n'.join(content_parts)
            print(f"✅ DOCX文件读取成功（{len(content)} 字符，包含 {len(doc.tables)} 个表格）")
            return content
        
        elif file_ext == '.pdf':
            # 读取PDF文件
            try:
                import fitz  # PyMuPDF
            except ImportError:
                raise ImportError("PyMuPDF未安装，无法读取PDF文件。请运行: pip install PyMuPDF")
            
            doc = fitz.open(file_path)
            content = ''
            page_count = len(doc)  # 先获取页数
            for page in doc:
                content += page.get_text()
            doc.close()
            
            print(f"✅ PDF文件读取成功（{len(content)} 字符，{page_count} 页）")
            return content
        
        else:
            raise ValueError(f"❌ 不支持的文件格式: {file_ext}\n   支持的格式：.txt, .docx, .pdf")
    
    except Exception as e:
        print(f"❌ 读取文件失败: {str(e)}")
        raise


def process_resume_and_match(resume_file: str,
                              output_dir: str = None,
                              top_k: int = 10,
                              use_llm: bool = False,
                              use_llm_for_profile: bool = False,
                              wechat_screenshots: Optional[List[str]] = None,
                              use_advanced_profile: bool = False,
                              output_tag: Optional[str] = None,
                              constraints: Optional["MatchConstraints"] = None) -> str:
    """
    处理简历文件并生成匹配报告（完整流程）
    
    Args:
        resume_file: 简历文件路径
        output_dir: 输出目录（默认使用config中的OUTPUT_DIR）
        top_k: 返回企业数量
        use_llm: 是否使用MiMo API生成报告
        use_llm_for_profile: 是否使用LLM生成人才画像摘要（基础版）
        wechat_screenshots: 微信聊天截图路径列表（PNG格式），用于高级人才画像
        use_advanced_profile: 是否使用高级人才画像（结合简历+微信截图）
        output_tag: 若设置（如时间戳），输出文件名为「简历名_匹配报告_{tag}.docx」，避免覆盖或目标文件被占用
        
    Returns:
        str: 生成的Word报告路径
    """
    print("=" * 60)
    print("处理简历并生成匹配报告")
    print("=" * 60)
    
    # ============ 第一步：读取简历 ============
    resume_text = read_resume_file(resume_file)
    
    # 显示简历摘要
    print(f"\n简历摘要（前300字符）:")
    print("-" * 60)
    print(resume_text[:300] + "..." if len(resume_text) > 300 else resume_text)
    
    # ============ 第二步：运行Pipeline ============
    if output_dir is None:
        output_dir = OUTPUT_DIR
    
    # 生成输出文件名（基于简历文件名）
    resume_name = Path(resume_file).stem
    tag_part = f"_{output_tag}" if output_tag else ""
    output_file = os.path.join(output_dir, f"{resume_name}_匹配报告{tag_part}.docx")
    
    print(f"\n📊 正在运行匹配Pipeline...")
    output_path = run_pipeline_and_export_word(
        resume_text=resume_text,
        top_k=top_k,
        output_file=output_file,
        use_llm=use_llm,
        use_llm_for_profile=use_llm_for_profile,  # 传递人才画像LLM生成参数
        wechat_screenshots=wechat_screenshots,  # 微信截图路径列表
        use_advanced_profile=use_advanced_profile,  # 是否使用高级人才画像
        constraints=constraints,
    )
    
    print("\n" + "=" * 60)
    print("✅ 简历处理完成！")
    print("=" * 60)
    print(f"输入文件: {resume_file}")
    print(f"输出报告: {output_path}")
    
    return output_path


if __name__ == "__main__":
    # 测试代码
    print("=" * 60)
    print("测试简历处理模块")
    print("=" * 60)
    
    # 创建测试简历文件（TXT格式）
    test_resume_path = "data/input/测试简历_半导体博士.txt"
    os.makedirs(os.path.dirname(test_resume_path), exist_ok=True)
    
    test_resume_content = """王明，男，1992年出生，博士学历，专业方向：集成电路、芯片设计。

教育背景：
- 2009-2013：复旦大学 微电子 学士
- 2013-2016：复旦大学 集成电路 硕士
- 2016-2020：清华大学 电子工程 博士

工作经历：
- 2020-2022：上海华虹集团 IC设计工程师
  - 参与28nm工艺芯片前端设计
  - 负责验证环境搭建与调试
  - 完成3个IP核设计项目

- 2022-至今：苏州纳芯微电子 高级IC设计工程师
  - 领导8人设计团队
  - 负责SoC架构设计与优化
  - 成功流片2款芯片（14nm、7nm）
  - 申请发明专利3项
  - 发表EI论文5篇

专业技能：
- 集成电路全流程设计（前端设计、验证、后端实现）
- 熟悉Verilog/SystemVerilog HDL
- 精通Synopsys EDA工具（VCS、DC、ICC2）
- 芯片架构设计与功耗优化
- 团队管理与技术攻关

求职意向：
- 期望地区：江苏省苏州市、无锡市、南京市
- 期望领域：半导体、集成电路、芯片设计
- 期望职位：IC设计经理/技术总监/首席架构师
"""
    
    with open(test_resume_path, 'w', encoding='utf-8') as f:
        f.write(test_resume_content)
    
    print(f"\n✅ 测试简历已创建: {test_resume_path}")
    
    # 处理测试简历
    try:
        output = process_resume_and_match(
            resume_file=test_resume_path,
            top_k=10,
            use_llm=False  # 不使用LLM（网络不可用）
        )
        
        print("\n" + "=" * 60)
        print("🎉 测试成功！")
        print("=" * 60)
        print(f"\n生成的报告: {output}")
        print("\n您可以：")
        print("1. 打开Word报告查看结果")
        print("2. 使用真实的简历文件测试（PDF、DOCX、TXT）")
        print("3. 批量处理多份简历")
        
    except Exception as e:
        print(f"\n❌ 测试失败: {str(e)}")
        import traceback
        traceback.print_exc()
