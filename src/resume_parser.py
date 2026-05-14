"""
简历解析模块
功能：解析PDF/Word简历，提取纯文本
"""
import os
import sys
from pathlib import Path
from typing import Optional

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF 未安装，PDF解析功能不可用")

try:
    from docx import Document
    PYDOCX_AVAILABLE = True
except ImportError:
    PYDOCX_AVAILABLE = False
    print("⚠️ python-docx 未安装，Word解析功能不可用")


def parse_pdf_resume(file_path: str) -> str:
    """
    解析PDF简历，返回纯文本
    
    Args:
        file_path: PDF文件路径
        
    Returns:
        str: 提取的纯文本
    """
    if not PYMUPDF_AVAILABLE:
        raise ImportError("PyMuPDF 未安装，请运行: pip install PyMuPDF")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    text_blocks = []
    
    try:
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                text_blocks.append(text.strip())
        
        doc.close()
        
    except Exception as e:
        raise RuntimeError(f"PDF解析失败: {e}")
    
    # 合并所有页面文本
    full_text = "\n\n".join(text_blocks)
    
    return full_text


def parse_word_resume(file_path: str) -> str:
    """
    解析Word简历，返回纯文本
    
    Args:
        file_path: Word文件路径（.docx）
        
    Returns:
        str: 提取的纯文本
    """
    if not PYDOCX_AVAILABLE:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    text_blocks = []
    
    try:
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_blocks.append(text)
        
        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_blocks.append(text)
                        
    except Exception as e:
        raise RuntimeError(f"Word解析失败: {e}")
    
    # 合并所有文本
    full_text = "\n".join(text_blocks)
    
    return full_text


def parse_resume(file_path: str) -> str:
    """
    自动识别文件类型并解析简历
    
    Args:
        file_path: 简历文件路径（支持.pdf, .docx）
        
    Returns:
        str: 提取的纯文本
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    ext = Path(file_path).suffix.lower()
    
    if ext == ".pdf":
        print(f"正在解析PDF简历: {file_path}")
        return parse_pdf_resume(file_path)
    elif ext in [".docx", ".doc"]:
        print(f"正在解析Word简历: {file_path}")
        return parse_word_resume(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}（仅支持 .pdf, .docx）")


def extract_resume_info(text: str) -> dict:
    """
    从简历文本中提取结构化信息（可选功能）
    
    Args:
        text: 简历纯文本
        
    Returns:
        dict: 结构化信息（姓名、教育背景、研究方向等）
    """
    info = {
        "姓名": "",
        "教育背景": [],
        "研究方向": [],
        "技能": [],
        "工作经历": []
    }
    
    lines = text.split("\n")
    
    # 简单的启发式提取（可以根据实际简历格式调整）
    for i, line in enumerate(lines):
        line_lower = line.lower()
        
        # 提取姓名（通常在前面几行）
        if i < 10 and not info["姓名"]:
            # 假设姓名行比较短且不包含常见关键词
            if 2 < len(line) < 20 and not any(kw in line_lower for kw in ["博士", "简历", "email", "电话"]):
                info["姓名"] = line.strip()
        
        # 提取教育背景
        if any(kw in line_lower for kw in ["博士", "硕士", "学士", "毕业", "学校", "university"]):
            info["教育背景"].append(line.strip())
            
        # 提取研究方向
        if any(kw in line_lower for kw in ["研究", "方向", "research", "方向", "领域"]):
            info["研究方向"].append(line.strip())
            
        # 提取技能
        if any(kw in line_lower for kw in ["技能", "熟悉", "精通", "掌握", "skill", "proficient"]):
            info["技能"].append(line.strip())
            
        # 提取工作经历
        if any(kw in line_lower for kw in ["工作", "经历", "公司", "任职", "experience", "work"]):
            info["工作经历"].append(line.strip())
    
    return info


if __name__ == "__main__":
    # 测试代码
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent))
    
    print("=" * 60)
    print("测试简历解析模块")
    print("=" * 60)
    
    # 测试Word简历解析
    test_word = "data/../资料/博士简历/word简历/WZ-00002701蔡程辉.docx"
    if os.path.exists(test_word):
        print(f"\n测试Word简历: {test_word}")
        text = parse_word_resume(test_word)
        print(f"✅ 解析成功，提取 {len(text)} 字符")
        print(f"\n前500字符:\n{text[:500]}...")
        
        # 测试信息提取
        info = extract_resume_info(text)
        print(f"\n提取的信息:")
        for k, v in info.items():
            if v:
                print(f"  {k}: {str(v)[:100]}...")
    else:
        print(f"\n⚠️ 测试文件不存在: {test_word}")
        print("  请确认文件路径是否正确")
